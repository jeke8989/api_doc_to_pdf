from fastapi import UploadFile
from typing import Dict
from io import BytesIO
from docx.document import Document as DocumentObject
from docx import Document
from docx2pdf import convert



# Данные для таблицы
add_rows = [
    ["1", "Description 1", "$567.43", "$567.43", "$567.43", "$567.43"],
    ["2", "Description 1", "$5.43", "$567.43", "$567.43", "$567.43"]
]

# Tестовый объект замены
replacements = {"<CCO Description>": "1111"}



#Получаем объект Document для работы
async def get_doc_file(file: UploadFile) -> DocumentObject:
    file_content = await file.read()
    file_stream = BytesIO(file_content)
    doc = Document(file_stream)
    return doc



# Заменяем переменные в файле
def replace_data(replacements: Dict, doc: DocumentObject, add_rows: list[list[str]]) -> str:
    tables = doc.tables
    # Обрабатываем все параграфы
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                para.text = para.text.replace(old_text, new_text)
    # Обрабатываем все таблицы
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for old_text, new_text in replacements.items():
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)
    # Добавляем ячейки
    table = tables[1]  # Предполагается, что вы хотите добавлять в таблицу с индексом 1
    for row in add_rows:
        new_row = table.add_row()  # Добавляем новую строку
        for index_cell, cell in enumerate(new_row.cells):
            if index_cell < len(row):  # Проверяем, что индекс ячейки не выходит за пределы
                cell.text = row[index_cell]  # Устанавливаем текст в ячейку
        
        
            
    
    
    doc.save('amendment3.docx')
    file_pdf_name = 'amendment3.pdf'
    convert('amendment3.docx',file_pdf_name)
    
    return file_pdf_name
